from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import FileResponse, HTMLResponse, StreamingResponse
from sqlalchemy.orm import Session
from datetime import date
from typing import Optional
import os
import io
from app import crud
from app.database import get_db
from app import models
from app.services.report_generator import generate_html_report, generate_pdf_report


# PDF imports
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# Register Cyrillic font
try:
    font_paths = [
        "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    ]
    for font_path in font_paths:
        if os.path.exists(font_path):
            pdfmetrics.registerFont(TTFont('CyrillicFont', font_path))
            break
except:
    pass

CYRILLIC_FONT = 'CyrillicFont'

router = APIRouter()

# Template files mapping (Latin names)
TEMPLATE_FILES = {
    'questionnaire': 'questionnaire.doc',
    'express': 'express_report.docx',
    'final': 'final_report.docx',
    'csv': 'ffp_report.docx',
    'ndt': 'ffp_report.docx',
    'epb': 'epb_report.docx',
}

# Project root - parent of backend folder
PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))


@router.get("/template/{report_type}")
async def download_template(report_type: str):
    """Download report template file"""
    if report_type not in TEMPLATE_FILES:
        raise HTTPException(status_code=404, detail="Template not found")
    
    filename = TEMPLATE_FILES[report_type]
    template_path = os.path.join(PROJECT_ROOT, filename)
    
    if not os.path.exists(template_path):
        raise HTTPException(status_code=404, detail=f"File not found: {filename}")
    
    return FileResponse(template_path, filename=filename)


@router.get("/filled/{report_type}")
async def generate_filled_report(
    report_type: str,
    db: Session = Depends(get_db)
):
    """Generate report filled with actual data - different structure per type"""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import tempfile
    from datetime import datetime
    
    if report_type not in TEMPLATE_FILES:
        raise HTTPException(status_code=404, detail="Report type not found")
    
    doc = Document()
    now = datetime.now()
    
    # Get data from database
    objects = db.query(models.Object).limit(20).all()
    inspections = db.query(models.Inspection).order_by(models.Inspection.date.desc()).limit(30).all()
    repair_works = db.query(models.RepairWork).order_by(models.RepairWork.created_at.desc()).limit(20).all()
    
    if report_type == 'questionnaire':
        # ОПРОСНЫЙ ЛИСТ - Survey form
        doc.add_heading('ОПРОСНЫЙ ЛИСТ', 0)
        doc.add_paragraph(f'Дата: {now.strftime("%d.%m.%Y")}')
        doc.add_paragraph()
        
        doc.add_heading('Характеристики трубопровода', level=1)
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Table Grid'
        rows = table.rows
        rows[0].cells[0].text = 'Наименование объекта'
        rows[0].cells[1].text = objects[0].object_name if objects else 'Нет данных'
        rows[1].cells[0].text = 'Тип объекта'
        rows[1].cells[1].text = objects[0].object_type if objects else 'Нет данных'
        rows[2].cells[0].text = 'Год постройки'
        rows[2].cells[1].text = str(objects[0].year) if objects and objects[0].year else 'Нет данных'
        rows[3].cells[0].text = 'Материал'
        rows[3].cells[1].text = objects[0].material if objects and objects[0].material else 'Нет данных'
        rows[4].cells[0].text = 'Количество объектов'
        rows[4].cells[1].text = str(len(objects))
        rows[5].cells[0].text = 'Количество обследований'
        rows[5].cells[1].text = str(len(inspections))
        
        doc.add_paragraph()
        doc.add_heading('Требования к диагностике', level=1)
        doc.add_paragraph('1. Визуальный и измерительный контроль (ВИК)')
        doc.add_paragraph('2. Ультразвуковая толщинометрия (UTWM)')
        doc.add_paragraph('3. Магнитопорошковый контроль (МПК)')
        
    elif report_type == 'express':
        # ЭКСПРЕСС-ОТЧЁТ - Quick summary
        doc.add_heading('ЭКСПРЕСС-ОТЧЁТ', 0)
        doc.add_paragraph(f'Дата формирования: {now.strftime("%d.%m.%Y %H:%M")}')
        doc.add_paragraph()
        
        doc.add_heading('Краткая сводка', level=1)
        defects_found = sum(1 for i in inspections if i.defect_found)
        doc.add_paragraph(f'Всего обследований: {len(inspections)}')
        doc.add_paragraph(f'Обнаружено дефектов: {defects_found}')
        doc.add_paragraph(f'Без дефектов: {len(inspections) - defects_found}')
        
        doc.add_heading('Последние обследования', level=1)
        if inspections:
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            header = table.rows[0].cells
            header[0].text = 'Дата'
            header[1].text = 'Метод'
            header[2].text = 'Дефект'
            header[3].text = 'Оценка'
            for insp in inspections[:10]:
                row = table.add_row().cells
                row[0].text = insp.date.strftime('%d.%m.%Y') if insp.date else '-'
                row[1].text = insp.method.value if insp.method else '-'
                row[2].text = 'Да' if insp.defect_found else 'Нет'
                row[3].text = insp.quality_grade.value if insp.quality_grade else '-'
                
    elif report_type == 'final':
        # ЗАКЛЮЧИТЕЛЬНЫЙ ОТЧЁТ - Full documentation
        doc.add_heading('ЗАКЛЮЧИТЕЛЬНЫЙ ОТЧЁТ', 0)
        doc.add_heading('по результатам диагностики', level=2)
        doc.add_paragraph(f'Дата: {now.strftime("%d.%m.%Y")}')
        doc.add_paragraph()
        
        doc.add_heading('1. Общие сведения', level=1)
        doc.add_paragraph(f'Количество обследованных объектов: {len(objects)}')
        doc.add_paragraph(f'Период диагностики: {inspections[-1].date.strftime("%d.%m.%Y") if inspections else "-"} - {inspections[0].date.strftime("%d.%m.%Y") if inspections else "-"}')
        
        doc.add_heading('2. Результаты диагностики', level=1)
        methods = {}
        for insp in inspections:
            m = insp.method.value if insp.method else 'Другой'
            methods[m] = methods.get(m, 0) + 1
        for method, count in methods.items():
            doc.add_paragraph(f'{method}: {count} обследований')
            
        doc.add_heading('3. Выявленные дефекты', level=1)
        defects = [i for i in inspections if i.defect_found]
        if defects:
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            header = table.rows[0].cells
            header[0].text = 'Дата'
            header[1].text = 'Метод'
            header[2].text = 'Описание'
            for d in defects[:15]:
                row = table.add_row().cells
                row[0].text = d.date.strftime('%d.%m.%Y') if d.date else '-'
                row[1].text = d.method.value if d.method else '-'
                row[2].text = d.defect_description or 'Нет описания'
        else:
            doc.add_paragraph('Дефекты не выявлены')
            
        doc.add_heading('4. Заключение', level=1)
        doc.add_paragraph('По результатам проведенной диагностики состояние объектов оценивается как удовлетворительное.')
        
    elif report_type in ['csv', 'ndt']:
        # CSV/FFP и ДДК-ОТЧЁТ - Data table format
        title = 'ОТЧЁТ ПО ДАННЫМ FFP' if report_type == 'csv' else 'ДДК-ОТЧЁТ (Верификация дефектов)'
        doc.add_heading(title, 0)
        doc.add_paragraph(f'Дата: {now.strftime("%d.%m.%Y")}')
        doc.add_paragraph()
        
        doc.add_heading('Данные обследований', level=1)
        if inspections:
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Table Grid'
            header = table.rows[0].cells
            header[0].text = 'ID'
            header[1].text = 'Объект'
            header[2].text = 'Дата'
            header[3].text = 'Метод'
            header[4].text = 'Дефект'
            header[5].text = 'Параметры'
            for insp in inspections:
                row = table.add_row().cells
                row[0].text = str(insp.diag_id)
                obj_name = insp.object.object_name if insp.object else '-'
                row[1].text = obj_name[:20] if len(obj_name) > 20 else obj_name
                row[2].text = insp.date.strftime('%d.%m.%Y') if insp.date else '-'
                row[3].text = insp.method.value if insp.method else '-'
                row[4].text = 'Да' if insp.defect_found else 'Нет'
                params = f'{insp.param1 or "-"}/{insp.param2 or "-"}/{insp.param3 or "-"}'
                row[5].text = params
                
    elif report_type == 'epb':
        # ЗАКЛЮЧЕНИЕ ЭПБ - Safety expertise
        doc.add_heading('ЗАКЛЮЧЕНИЕ', 0)
        doc.add_heading('экспертизы промышленной безопасности', level=2)
        doc.add_paragraph(f'Дата: {now.strftime("%d.%m.%Y")}')
        doc.add_paragraph()
        
        doc.add_heading('1. Объект экспертизы', level=1)
        doc.add_paragraph(f'Наименование: {objects[0].object_name if objects else "Трубопровод"}')
        doc.add_paragraph(f'Тип: Магистральный трубопровод')
        
        doc.add_heading('2. Цель экспертизы', level=1)
        doc.add_paragraph('Определение соответствия объекта требованиям промышленной безопасности.')
        
        doc.add_heading('3. Результаты экспертизы', level=1)
        doc.add_paragraph(f'Проведено обследований: {len(inspections)}')
        defects = sum(1 for i in inspections if i.defect_found)
        doc.add_paragraph(f'Выявлено дефектов: {defects}')
        
        doc.add_heading('4. Выводы', level=1)
        if defects > len(inspections) * 0.3:
            doc.add_paragraph('Объект требует проведения ремонтных работ.')
        else:
            doc.add_paragraph('Объект соответствует требованиям промышленной безопасности.')
            
        doc.add_heading('5. Рекомендации', level=1)
        doc.add_paragraph('1. Продолжить плановый мониторинг состояния')
        doc.add_paragraph('2. Устранить выявленные дефекты в установленные сроки')
        doc.add_paragraph('3. Провести повторную диагностику через 12 месяцев')
    
    # Save document
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    temp_file.close()
    
    report_names = {
        'questionnaire': 'oprosny_list',
        'express': 'express_report', 
        'final': 'final_report',
        'csv': 'ffp_data',
        'ndt': 'ddk_report',
        'epb': 'epb_conclusion'
    }
    
    filename = f'{report_names.get(report_type, "report")}_{now.strftime("%Y%m%d_%H%M%S")}.docx'
    
    return FileResponse(
        temp_file.name,
        filename=filename,
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )



@router.get("/generate")
async def generate_report(
    format: str = Query("html", regex="^(html|pdf)$"),
    date_from: Optional[date] = None,
    date_to: Optional[date] = None,
    pipeline_id: Optional[str] = None,
    risk_level: Optional[str] = None,
    db: Session = Depends(get_db)
):
    """Generate a report in HTML or PDF format"""
    
    try:
        # Get data for report
        inspections = crud.get_inspections(
            db,
            date_from=date_from,
            date_to=date_to,
            pipeline_id=pipeline_id,
            risk_level=risk_level,
            limit=10000
        )
        
        # Eagerly load object relationships for excavation recommendations
        from sqlalchemy.orm import joinedload
        inspections_with_objects = db.query(models.Inspection).options(
            joinedload(models.Inspection.object)
        ).filter(models.Inspection.id.in_([i.id for i in inspections])).all()
        
        stats = crud.get_dashboard_stats(db)
        
        if format == "html":
            html_content = generate_html_report(inspections_with_objects, stats, date_from, date_to)
            return HTMLResponse(content=html_content)
        else:
            pdf_path = generate_pdf_report(inspections_with_objects, stats, date_from, date_to)
            return FileResponse(
                pdf_path,
                media_type="application/pdf",
                filename="integrityos_report.pdf"
            )
            
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Report generation failed: {str(e)}")


@router.get("/defect/{work_id}/pdf")
async def generate_defect_report(
    work_id: int,
    db: Session = Depends(get_db)
):
    """Generate a PDF report for a specific repair work with photos"""
    
    # Get work with related inspection
    work = db.query(models.RepairWork).filter(models.RepairWork.id == work_id).first()
    if not work:
        raise HTTPException(status_code=404, detail="Work not found")
    
    inspection = db.query(models.Inspection).filter(
        models.Inspection.id == work.inspection_id
    ).first()
    
    # Get media files
    media_items = db.query(models.Media).filter(
        models.Media.inspection_id == work.inspection_id
    ).all()
    
    # Create PDF in memory
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=2*cm, leftMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)
    
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('title', parent=styles['Heading1'], fontName=CYRILLIC_FONT, alignment=TA_CENTER, spaceAfter=20)
    heading_style = ParagraphStyle('heading', parent=styles['Heading2'], fontName=CYRILLIC_FONT, spaceAfter=10, spaceBefore=15)
    normal_style = ParagraphStyle('normal', parent=styles['Normal'], fontName=CYRILLIC_FONT)
    
    elements = []
    
    # Title
    elements.append(Paragraph("Акт выполненных работ", title_style))
    elements.append(Spacer(1, 0.5*cm))
    
    # Work info table
    status_labels = {
        'planned': 'Запланировано',
        'in_progress': 'В работе',
        'pending_approval': 'Ожидает проверки',
        'completed': 'Завершено',
        'cancelled': 'Отменено'
    }
    
    priority_labels = {
        'low': 'Низкий',
        'medium': 'Средний', 
        'high': 'Высокий',
        'critical': 'Критический'
    }
    
    work_data = [
        ['Наименование работы:', work.title],
        ['Описание:', work.description or '-'],
        ['Приоритет:', priority_labels.get(work.priority, work.priority)],
        ['Статус:', status_labels.get(work.status.value if hasattr(work.status, 'value') else work.status, str(work.status))],
        ['Плановая дата:', str(work.planned_date) if work.planned_date else '-'],
        ['Дата завершения:', str(work.completed_date) if work.completed_date else '-'],
        ['Дата создания:', str(work.created_at.date()) if work.created_at else '-'],
    ]
    
    work_table = Table(work_data, colWidths=[5*cm, 11*cm])
    work_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (0, -1), 'RIGHT'),
        ('ALIGN', (1, 0), (1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, -1), CYRILLIC_FONT),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOX', (0, 0), (-1, -1), 1, colors.black),
        ('INNERGRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('LEFTPADDING', (0, 0), (-1, -1), 8),
        ('RIGHTPADDING', (0, 0), (-1, -1), 8),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
    ]))
    elements.append(work_table)
    elements.append(Spacer(1, 0.5*cm))
    
    # Inspection info
    if inspection:
        elements.append(Paragraph("Данные обследования", heading_style))
        object_name = inspection.object.object_name if inspection.object else '-'
        method_str = inspection.method.value if hasattr(inspection.method, 'value') else str(inspection.method) if inspection.method else '-'
        quality_str = inspection.quality_grade.value if hasattr(inspection.quality_grade, 'value') else str(inspection.quality_grade) if inspection.quality_grade else '-'
        insp_data = [
            ['Объект:', object_name],
            ['Дата обследования:', str(inspection.date) if inspection.date else '-'],
            ['Метод:', method_str],
            ['Оценка качества:', quality_str],
        ]
        
        insp_table = Table(insp_data, colWidths=[5*cm, 11*cm])
        insp_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
            ('ALIGN', (0, 0), (0, -1), 'RIGHT'),
            ('ALIGN', (1, 0), (1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, -1), CYRILLIC_FONT),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOX', (0, 0), (-1, -1), 1, colors.black),
            ('INNERGRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('LEFTPADDING', (0, 0), (-1, -1), 8),
            ('RIGHTPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ]))
        elements.append(insp_table)
        elements.append(Spacer(1, 0.5*cm))
    
    # Photos section
    before_photos = [m for m in media_items if m.is_before]
    after_photos = [m for m in media_items if not m.is_before]
    
    media_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "media")
    
    if before_photos:
        elements.append(Paragraph("Фото ДО ремонта", heading_style))
        for photo in before_photos:
            photo_path = os.path.join(media_dir, photo.filename)
            if os.path.exists(photo_path):
                try:
                    img = Image(photo_path, width=12*cm, height=9*cm)
                    img.hAlign = 'CENTER'
                    elements.append(img)
                    elements.append(Spacer(1, 0.3*cm))
                except Exception as e:
                    elements.append(Paragraph(f"[Фото недоступно: {photo.filename}]", normal_style))
    
    if after_photos:
        elements.append(Paragraph("Фото ПОСЛЕ ремонта", heading_style))
        for photo in after_photos:
            photo_path = os.path.join(media_dir, photo.filename)
            if os.path.exists(photo_path):
                try:
                    img = Image(photo_path, width=12*cm, height=9*cm)
                    img.hAlign = 'CENTER'
                    elements.append(img)
                    elements.append(Spacer(1, 0.3*cm))
                except Exception as e:
                    elements.append(Paragraph(f"[Фото недоступно: {photo.filename}]", normal_style))
    
    # Notes
    if work.notes:
        elements.append(Paragraph("Примечания", heading_style))
        elements.append(Paragraph(work.notes.replace('\n', '<br/>'), normal_style))
    
    # Build PDF
    doc.build(elements)
    buffer.seek(0)
    
    filename = f"defect_report_{work_id}.pdf"
    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

